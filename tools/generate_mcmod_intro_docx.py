from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "NetMusicCanPlayBili-MCMOD介绍正文.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(document: Document, text: str, *, level: int = 0) -> None:
    paragraph = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(47, 54, 64)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(7)

    for name, size, color in (
        ("Title", 24, "1F4E78"),
        ("Subtitle", 11, "607D8B"),
        ("Heading 1", 16, "1F4E78"),
        ("Heading 2", 13, "2F75B5"),
    ):
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if name != "Subtitle":
            style.font.bold = True

    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)


def add_info_table(document: Document) -> None:
    rows = (
        ("模组名称", "Net Music Can Play Bili"),
        ("当前版本", "0.5.5-beta"),
        ("适用版本", "Minecraft 26.1.2 / NeoForge 26.x / Java 25"),
        ("前置模组", "NetMusic 1.5.1 或更高版本（必装）"),
        ("可选兼容", "Curios 12 或更高版本"),
        ("运行环境", "客户端与服务端均需安装"),
        ("作者", "zhongbai233"),
        ("开源许可", "MIT License"),
    )
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        set_cell_text(table.cell(index, 0), label, bold=True, color="FFFFFF")
        set_cell_shading(table.cell(index, 0), "2F75B5")
        set_cell_text(table.cell(index, 1), value)
        if index % 2 == 1:
            set_cell_shading(table.cell(index, 1), "EAF2F8")


def add_note(document: Document, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF2CC")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    title_run = paragraph.add_run(f"{title}：")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(156, 101, 0)
    paragraph.add_run(text)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    configure_styles(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Net Music Can Play Bili")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("MCMOD 模组介绍正文｜把 Bilibili 影音带进 Minecraft")

    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    opening = lead.add_run("不只是把一首歌塞进唱片里，而是把一整套影音系统搬进方块世界。")
    opening.bold = True
    opening.font.color.rgb = RGBColor(31, 78, 120)
    lead.add_run(
        "Net Music Can Play Bili 是 NetMusic 的 NeoForge 附属模组，为其加入 Bilibili 视频与直播音频的解析、刻录和播放支持。"
        "玩家可以将 BV 号、av 号或视频链接刻录成 NetMusic 唱片，再通过现代化唱片机、投影仪、音响、MP4、Pad、耳机和全息眼镜，"
        "搭建从客厅影院到冒险配乐地图的完整游戏内媒体系统。"
    )

    document.add_heading("基本信息", level=1)
    add_info_table(document)

    document.add_heading("核心玩法", level=1)

    document.add_heading("把 Bilibili 视频刻成唱片", level=2)
    document.add_paragraph(
        "在 NetMusic 电脑中登录 Bilibili 后，将 BV 号、av 号或含有对应编号的视频链接填入 CD 刻录机，即可制作可播放的 NetMusic 音乐 CD。"
        "多 P 视频可以指定分 P 序号；登录后还可提高部分视频及高画质资源的解析成功率，并可选择是否优先尝试 Dolby 音频。"
    )
    document.add_paragraph(
        "如果你更喜欢实时内容，NetMusic 的广播喇叭还可输入“live:直播间ID”，直接播放对应 Bilibili 直播间的音频。"
    )

    document.add_heading("现代化唱片机与自动化", level=2)
    document.add_paragraph(
        "现代化唱片机提供播放、暂停、重播、单曲循环、进度拖动和音量调节等完整控制。它支持漏斗与 NeoForge 物品传输接口，"
        "可自动插入或取出唱片，并提供“播完提取”和“自由提取”两种自动化策略。"
    )
    document.add_paragraph("唱片机还拥有四种红石模式：高信号播放、低信号播放、脉冲切换和忽略红石。比较器可以输出当前曲目进度，因此你可以制作：")
    add_bullet(document, "随音乐推进逐格亮起的进度灯带；")
    add_bullet(document, "曲目结束后自动换片的红石装置；")
    add_bullet(document, "由按钮、压力板或机关控制的场景配乐系统。")

    document.add_heading("搭建属于你的影音空间", level=2)
    document.add_paragraph("唱片机可与多种影音设备绑定，设备放置位置更加自由，适合打造影院、舞台、展厅或私人音乐室。")
    add_bullet(document, "歌词投影仪：可调整朝向、俯仰、偏移和缩放，支持多种歌词显示模式及 AI 字幕。")
    add_bullet(document, "视频投影仪：支持从 360P 到 8K 的多档画质，并可细调投影位置、角度和大小。")
    add_bullet(document, "音响：可分配 7.1.4 声道位置，支持独立静音、音量调节与未分配声道融合，组合摆放可获得更鲜明的空间音频体验。")

    document.add_heading("MP4：随身携带的播放列表", level=2)
    document.add_paragraph(
        "MP4 最多可装入 18 张 NetMusic 唱片。手持右键即可进入聚焦界面，管理播放列表、切换歌曲、播放视频、选择画质与字幕，并调整音量和循环状态。"
        "视频只会在设备位于快捷栏或副手等活跃位置时持续解码，以减少不必要的资源占用；切换物品后，音频仍可继续播放。"
    )

    document.add_heading("Pad：把音乐写进地图", level=2)
    document.add_paragraph(
        "Pad 并非放大版 MP4，而是一套以地图点位为核心的场景媒体系统。你可以把唱片拖到地图上创建音乐 Pin，设置触发半径、播放方式、循环和音量，"
        "也可以隐藏点位，让玩家在毫无提示的情况下踏入一段专属配乐。"
    )
    document.add_paragraph("Pad 特别适合制作：")
    add_bullet(document, "景区或服务器地标的语音导览；")
    add_bullet(document, "RPG 地图的区域主题曲与剧情配乐；")
    add_bullet(document, "密室解谜中的声音线索；")
    add_bullet(document, "只有走进特定角落才能发现的音乐彩蛋。")
    document.add_paragraph(
        "编辑完成后可以发布锁定副本供其他玩家使用。草稿仍可继续修改，并会同步到对应副本。单个 Pad 最多保存 64 个媒体和 128 个点位。"
    )

    document.add_heading("耳机与全息眼镜", level=2)
    document.add_paragraph(
        "通过媒体管理工具，可以把唱片机、MP4 或 Pad 与穿戴设备连接。耳机会将设备声音私有地传给佩戴者；全息眼镜则能将绑定媒体显示在 HUD 全息层中，"
        "最多同时绑定 4 个媒体源。安装 Curios 后，这些设备可以放入饰品头部槽；未安装时则使用原版头盔栏。"
    )

    document.add_heading("多人服务器支持", level=1)
    document.add_paragraph(
        "模组不仅考虑了播放，也提供面向服务器的内容管理能力。管理员可以启用视频链接白名单、查看活跃媒体源，并接收玩家通过媒体管理工具提交的音源举报。"
        "重复举报会自动合并，即使当时没有管理员在线，记录也会保存，方便后续审核。"
    )
    add_bullet(document, "支持限制玩家可播放的视频链接；")
    add_bullet(document, "支持审计附近或服务器中的活跃媒体源；")
    add_bullet(document, "设备配置会校验距离、建造权限与玩家权限；")
    add_bullet(document, "适合生存服、城建服、剧情服和公开服务器使用。")

    document.add_heading("五分钟快速上手", level=1)
    add_number(document, "安装本模组与前置 NetMusic；多人游戏中客户端和服务端均需安装，且版本保持一致。")
    add_number(document, "按 NetMusic 的玩法制作电脑和 CD 刻录机，在电脑中点击“B站登录”并扫码。")
    add_number(document, "把 BV 号、av 号或 Bilibili 视频链接填入 CD 刻录机；多 P 视频按需填写分 P 序号。")
    add_number(document, "将刻录完成的唱片插入现代化唱片机，即可开始播放。")
    add_number(document, "手持歌词投影仪、视频投影仪或音响右键唱片机完成预绑定，再把设备放到合适位置。")

    document.add_heading("安装与兼容说明", level=1)
    add_bullet(document, "当前项目版本为 beta，更新前建议备份世界与配置。")
    add_bullet(document, "NetMusic 为必装前置；Curios 为可选兼容，不安装也可使用穿戴设备。")
    add_bullet(document, "模组已打包常见平台所需的原生解码库，普通玩家通常无须另行安装系统 FFmpeg。")
    add_bullet(document, "高画质能否取得取决于视频源、Bilibili 登录状态、账号权限和网络环境。画面卡顿时建议优先降低到 720P 或更低画质。")
    add_bullet(document, "Bilibili 登录信息保存在本地配置中，请勿公开上传或转发含 Cookie 的配置文件。")

    add_note(
        document,
        "内容与版权提醒",
        "本模组只提供媒体解析与游戏内播放能力，不附带任何 Bilibili 视频或音频内容。请遵守相关平台规则、著作权法律及服务器规定，"
        "仅播放你有权访问和使用的内容。公开服务器建议启用链接白名单与举报审计功能。",
    )

    document.add_heading("为什么值得一试？", level=1)
    document.add_paragraph(
        "如果你只想在挖矿时听一首歌，MP4 和耳机可以满足你；如果你想为基地搭建家庭影院，投影仪与空间音响已经准备就绪；"
        "如果你正在制作 RPG、景区或解谜地图，Pad 的位置触发媒体又能成为新的叙事工具。Net Music Can Play Bili 将 Bilibili 播放能力真正融入了方块、物品、红石、地图和多人管理，"
        "让影音内容不再只是悬浮在游戏外的一块播放器窗口，而成为 Minecraft 世界中可以摆放、携带、连接和设计的一部分。"
    )

    ending = document.add_paragraph()
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ending.add_run("刻下一张唱片，接好投影与音响——你的方块世界，现在可以开演了。")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 120)

    document.add_paragraph()
    source = document.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source_run = source.add_run("正文依据项目 README、游玩指南及模组元数据整理｜生成日期：2026-07-23")
    source_run.italic = True
    source_run.font.size = Pt(8.5)
    source_run.font.color.rgb = RGBColor(127, 127, 127)
    return document


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(OUTPUT)